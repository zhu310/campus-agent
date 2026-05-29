"""文件文本解析服务。

支持 txt/md/docx/pdf 的文本提取，其中 docx 会额外兜底读取 XML，提升模板类
文档的解析成功率。
"""

from pathlib import Path
import re
from zipfile import BadZipFile, ZipFile
from xml.etree import ElementTree as ET

try:
    import pdfplumber
except Exception:  # pragma: no cover - optional runtime dependency
    pdfplumber = None

try:
    from docx import Document as DocxDocument
except Exception:  # pragma: no cover - optional runtime dependency
    DocxDocument = None


CHECKED_MARKS = "☑√✓■●"
UNCHECKED_MARKS = "□☐○"


def _clean_docx_cell(value: str) -> str:
    value = str(value or "").replace("\r", "\n").strip()
    value = re.sub(r"\n{2,}", "\n", value)
    value = re.sub(r"[ \t]{2,}", " ", value)
    return value.strip()


def _collapse_repeated_cells(cells: list[str]) -> list[str]:
    collapsed: list[str] = []
    for cell in cells:
        cleaned = _clean_docx_cell(cell)
        if not cleaned:
            continue
        if collapsed and collapsed[-1] == cleaned:
            continue
        collapsed.append(cleaned)
    return collapsed


def _looks_like_form_label(value: str) -> bool:
    compact = re.sub(r"\s+", "", value or "")
    if not compact or len(compact) > 32:
        return False
    if any(mark in compact for mark in CHECKED_MARKS + UNCHECKED_MARKS):
        return False
    if re.search(r"[。！？!?；;,，]", compact):
        return False
    return bool(re.search(r"[\u4e00-\u9fffA-Za-z]", compact))


def _looks_like_section_label(value: str) -> bool:
    compact = re.sub(r"\s+", "", value or "")
    return bool(re.search(r"(情况|成绩|理由|意见|审核|审批|信息)$", compact)) and len(compact) <= 12


def _selected_checkbox_value(value: str) -> str:
    text = re.sub(r"\s+", " ", value or "")
    for mark in CHECKED_MARKS:
        matches = re.findall(rf"([\u4e00-\u9fffA-Za-z0-9_./（）() -]{{1,30}})\s*{re.escape(mark)}", text)
        if matches:
            return matches[-1].strip(" :：/|")
    return ""


def _strip_leading_section_cell(cells: list[str]) -> list[str]:
    if len(cells) > 2 and _looks_like_section_label(cells[0]):
        return cells[1:]
    return cells


def _docx_table_kv_lines(cells: list[str]) -> list[str]:
    cells = _strip_leading_section_cell(cells)
    lines: list[str] = []
    seen: set[tuple[str, str]] = set()

    def add(label: str, value: str) -> None:
        label = _clean_docx_cell(label).strip(":：| ")
        value = _clean_docx_cell(value).strip(":：| ")
        selected = _selected_checkbox_value(value)
        if selected:
            value = selected
        if not label or not value or label == value:
            return
        key = (label, value)
        if key not in seen:
            seen.add(key)
            lines.append(f"{label}: {value}")

    for cell in cells:
        for match in re.finditer(r"([^:：|]{2,40})\s*[:：]\s*([^|]+)", cell):
            add(match.group(1), match.group(2))

    for index, cell in enumerate(cells):
        selected = _selected_checkbox_value(cell)
        if selected and index > 0 and ":" not in cell and "：" not in cell and _looks_like_form_label(cells[index - 1]):
            add(cells[index - 1], selected)

    for index, cell in enumerate(cells[:-1]):
        if not _looks_like_form_label(cell):
            continue
        value = cells[index + 1]
        if _looks_like_form_label(value):
            continue
        add(cell, value)
    return lines


def _extract_docx_xml_text(file_path: str) -> str:
    texts: list[str] = []
    namespaces = {
        "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
    }
    try:
        with ZipFile(file_path) as archive:
            names = [
                name
                for name in archive.namelist()
                if name.startswith("word/")
                and name.endswith(".xml")
                and (
                    name == "word/document.xml"
                    or name.startswith("word/header")
                    or name.startswith("word/footer")
                    or name.startswith("word/footnotes")
                    or name.startswith("word/endnotes")
                )
            ]
            for name in names:
                root = ET.fromstring(archive.read(name))
                for node in root.findall(".//w:t", namespaces):
                    if node.text and node.text.strip():
                        texts.append(node.text.strip())
                for node in root.findall(".//w:tab", namespaces):
                    texts.append("\t")
                for node in root.findall(".//w:br", namespaces):
                    texts.append("\n")
    except (BadZipFile, ET.ParseError, KeyError):
        return ""
    return "\n".join(texts)


def extract_text_from_file(file_path: str) -> str:
    path = Path(file_path)
    suffix = path.suffix.lower()

    if suffix in {'.txt', '.md'}:
        return path.read_text(encoding='utf-8', errors='ignore')

    if suffix == '.pdf':
        if pdfplumber is None:
            raise ValueError('PDF parsing requires pdfplumber. Please install backend requirements.')
        texts = []
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text() or ''
                if page_text.strip():
                    texts.append(page_text)
                for table in page.extract_tables() or []:
                    for row in table or []:
                        cells = [str(cell or '').strip() for cell in row if str(cell or '').strip()]
                        if cells:
                            texts.append(' | '.join(cells))
        return '\n'.join(texts)

    if suffix == '.docx':
        texts = []
        if DocxDocument is not None:
            doc = DocxDocument(file_path)
            paragraphs = [paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()]
            table_lines = []
            for table in doc.tables:
                for row in table.rows:
                    cells = _collapse_repeated_cells([cell.text for cell in row.cells])
                    if cells:
                        table_lines.append(' | '.join(cells))
                        table_lines.extend(_docx_table_kv_lines(cells))
            texts.append('\n'.join(paragraphs + table_lines))
        if not texts or len("\n".join(texts).strip()) < 80:
            texts.append(_extract_docx_xml_text(file_path))
        merged_lines: list[str] = []
        seen_lines: set[str] = set()
        for block in texts:
            for line in block.splitlines():
                line = line.strip()
                if line and line not in seen_lines:
                    seen_lines.add(line)
                    merged_lines.append(line)
        return "\n".join(merged_lines)

    raise ValueError(f'Unsupported file type: {suffix}')
